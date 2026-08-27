"""本地范文文档读取与拆分。

支持 .docx（python-docx）/ .txt（UTF-8/GBK 自动检测）/ .xml（剥标签后读）。
按「范文N话题：」标记把每个文件拆分成独立单篇，供每日轮转解析。
"""
from __future__ import annotations

import logging
import os
import re

log = logging.getLogger(__name__)

# 范文标记：范文一话题：XXX / 范文1话题：XXX
_FANWEN_RE = re.compile(r"^范文([一二三四五六七八九十百\d]+)话题[:：]\s*(.*)$")
# 用于回退的通用标题：含"话题"或"：xxx"的短行
_FALLBACK_RE = re.compile(r"^(?:第?.+篇|话题[:：]?.{2,30})$")

MIN_FANWEN_LEN = 300  # 短于此长度的篇自动跳过


def read_file_text(path: str) -> str:
    """读取 docx/txt/xml 为纯文本。"""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".txt":
        return _read_txt(path)
    if ext in (".xml", ".docm"):
        return _read_xml(path)
    raise ValueError(f"不支持的文件类型: {ext}（仅支持 .docx/.txt/.xml）")


def _read_docx(path: str) -> str:
    from docx import Document
    d = Document(path)
    paras = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    # 表格内容也并入（范文可能放在表格里）
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    paras.append(t)
    return "\n".join(paras)


def _read_txt(path: str) -> str:
    for enc in ("utf-8", "gbk", "utf-16"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别文本编码: {path}")


def _read_xml(path: str) -> str:
    from bs4 import BeautifulSoup
    with open(path, encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "xml")
        return " ".join(soup.get_text(" ", strip=True).split())


def split_fanwen(text: str, source_file: str = "") -> list[dict]:
    """按「范文N话题：」标记拆分文本为单篇列表。

    返回：[{title, content, index, file}]；无标记时整篇视为一篇。
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    articles: list[dict] = []
    cur: dict | None = None

    def flush():
        nonlocal cur
        if cur and cur["content"].strip():
            articles.append(cur)
        cur = None

    for ln in lines:
        m = _FANWEN_RE.match(ln)
        if m:
            flush()
            cur = {"title": ln, "index": m.group(1), "file": source_file, "content": ""}
            topic = m.group(2).strip()
            if topic:
                cur["topic"] = topic
            continue
        # 连续标题行（正文还没开始）：若当前篇内容为空且这是第二行标题，追加到标题
        if cur is None:
            cur = {"title": ln, "index": str(len(articles) + 1), "file": source_file,
                   "topic": ln, "content": ""}
            continue
        cur["content"] += ln + "\n"

    flush()

    # 无标记时整篇作为一篇
    if not articles:
        articles = [{"title": os.path.basename(source_file) or "范文",
                     "index": "1", "file": source_file,
                     "topic": os.path.splitext(os.path.basename(source_file))[0],
                     "content": text.strip()}]
    return articles


def scan_sucai(sucai_dir: str) -> list[dict]:
    """扫描 sucai 目录下所有支持的文件，返回拆分后的范文列表（按文件名排序）。"""
    results: list[dict] = []
    if not os.path.isdir(sucai_dir):
        return results
    for fn in sorted(os.listdir(sucai_dir)):
        if fn.startswith("~$"):
            continue  # Word 临时锁文件
        path = os.path.join(sucai_dir, fn)
        if not os.path.isfile(path):
            continue
        ext = os.path.splitext(fn)[1].lower()
        if ext not in (".docx", ".txt", ".xml", ".docm"):
            continue
        try:
            text = read_file_text(path)
            for art in split_fanwen(text, source_file=fn):
                art["path"] = path
                art["mtime"] = os.path.getmtime(path)
                results.append(art)
        except Exception as e:  # noqa: BLE001
            log.warning("解析文件失败 %s: %s", fn, e)
    return results
