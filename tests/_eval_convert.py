# -*- coding: utf-8 -*-
"""评估：docx 转 txt/md 的实际收益（体积/内容质量对比）。"""
import sys, io, os
sys.path.insert(0, ".")
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from app import docreader

total_docx = 0
total_txt = 0
for fn in sorted(os.listdir("sucai")):
    if not fn.endswith(".docx") or fn.startswith("~$"):
        continue
    path = os.path.join("sucai", fn)
    size = os.path.getsize(path)
    total_docx += size
    text = docreader.read_file_text(path)
    txt_size = len(text.encode("utf-8"))
    total_txt += txt_size
    arts = docreader.split_fanwen(text, source_file=fn)
    print(f"{fn}: docx {size//1024}KB -> txt {txt_size//1024}KB | {len(arts)}篇")

print(f"\n总计: docx {total_docx//1024}KB -> txt {total_txt//1024}KB ({(1-total_txt/total_docx)*100:.0f}% 缩小)")

# 检查转换是否有内容损失（docx 表格等）
print("\n=== 表格检查（docx 中是否有表格内容）===")
try:
    from docx import Document
    has_table = 0
    for fn in os.listdir("sucai"):
        if not fn.endswith(".docx") or fn.startswith("~$"):
            continue
        d = Document(os.path.join("sucai", fn))
        if d.tables:
            has_table += 1
            print(f"  {fn}: {len(d.tables)} 个表格")
    if not has_table:
        print("  无表格（纯文本结构，转换无损失）")
except Exception as e:
    print("检查失败:", e)
